import os
from pathlib import Path
from typing import List, Dict
import PyPDF2
from docx import Document

from config import settings

def load_pdf(file_path: str) -> List[str]:
    """Загрузка текста из PDF"""
    text_chunks = []
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_chunks.append(text)
    except Exception as e:
        print(f"Ошибка при загрузке PDF {file_path}: {e}")
    return text_chunks

def load_docx(file_path: str) -> List[str]:
    """Загрузка текста из DOCX"""
    text_chunks = []
    try:
        # Пробуем открыть документ
        try:
            doc = Document(file_path)
        except Exception as e:
            # Если не удалось открыть стандартным способом, пробуем альтернативный
            print(f"Предупреждение при открытии DOCX {file_path}: {e}")
            # Пробуем через zipfile напрямую
            import zipfile
            import xml.etree.ElementTree as ET
            
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    # Ищем основной документ
                    if 'word/document.xml' in zip_ref.namelist():
                        xml_content = zip_ref.read('word/document.xml')
                        root = ET.fromstring(xml_content)
                        # Простое извлечение текста из XML
                        full_text = []
                        for elem in root.iter():
                            if elem.text:
                                full_text.append(elem.text)
                        doc_text = ' '.join(full_text)
                        if doc_text.strip():
                            # Разбиваем на чанки
                            chunk_size = 1000
                            return [doc_text[i:i+chunk_size] for i in range(0, len(doc_text), chunk_size)]
            except Exception as e2:
                print(f"Ошибка при альтернативной загрузке DOCX {file_path}: {e2}")
                return []
        
        full_text = []
        # Извлекаем текст из параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(' | '.join(row_text))
        
        # Разбиваем на чанки по 1000 символов
        chunk = ""
        for text in full_text:
            if len(chunk) + len(text) > 1000:
                if chunk:
                    text_chunks.append(chunk)
                chunk = text
            else:
                chunk += "\n" + text if chunk else text
        
        if chunk:
            text_chunks.append(chunk)
    except Exception as e:
        print(f"Ошибка при загрузке DOCX {file_path}: {e}")
        import traceback
        traceback.print_exc()
    return text_chunks

def load_text_file(file_path: str) -> List[str]:
    """Загрузка текста из текстового файла"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            # Разбиваем на чанки
            chunk_size = 1000
            return [content[i:i+chunk_size] for i in range(0, len(content), chunk_size)]
    except Exception as e:
        print(f"Ошибка при загрузке текстового файла {file_path}: {e}")
        return []

def _is_vnd_upload_path(file_path: str) -> bool:
    """Файл из папки загрузки ВНД для анализа (IN и подпапки)."""
    if not file_path:
        return False
    try:
        in_root = Path(settings.in_folder).resolve()
        candidate = Path(file_path)
        if not candidate.is_absolute():
            candidate = in_root / candidate
        resolved = candidate.resolve()
        if resolved.is_relative_to(in_root):
            return True
        same_name = in_root / resolved.name
        return same_name.is_file() and os.path.samefile(resolved, same_name)
    except (ValueError, OSError, TypeError):
        return False


def extract_full_text(file_path: str, *, apply_vnd_mask: bool | None = None) -> str:
    """Извлечь полный текст из документа (PDF, DOCX, TXT)."""
    file_path_obj = Path(file_path)
    if not file_path_obj.exists():
        return ""

    file_ext = file_path_obj.suffix.lower()
    chunks: List[str] = []

    if file_ext == ".pdf":
        chunks = load_pdf(str(file_path))
    elif file_ext in [".docx", ".doc"]:
        chunks = load_docx(str(file_path))
    elif file_ext in [".txt", ".md"]:
        chunks = load_text_file(str(file_path))

    text = "\n".join(chunks)
    should_mask = apply_vnd_mask if apply_vnd_mask is not None else _is_vnd_upload_path(str(file_path))
    if text and should_mask:
        from vnd_masking import mask_vnd_sensitive_data

        text = mask_vnd_sensitive_data(text)
    return text


def process_single_file(file_path: str, store, collection_name: str) -> Dict:
    """Обработка одного файла"""
    file_path_obj = Path(file_path)
    
    if not file_path_obj.exists():
        error_msg = f"Файл {file_path} не существует"
        print(error_msg)
        return {"status": "error", "message": error_msg}
    
    if not file_path_obj.is_file():
        error_msg = f"{file_path} не является файлом"
        print(error_msg)
        return {"status": "error", "message": error_msg}
    
    print(f"Обработка файла: {file_path}")
    file_ext = file_path_obj.suffix.lower()
    
    # Поддерживаемые расширения
    supported_extensions = {".pdf", ".docx", ".doc", ".txt", ".md"}
    
    if not file_ext or file_ext not in supported_extensions:
        error_msg = f"Неподдерживаемый формат файла: {file_ext}"
        print(error_msg)
        return {"status": "error", "message": error_msg}
    
    chunks = []
    
    try:
        if file_ext == ".pdf":
            chunks = load_pdf(str(file_path))
        elif file_ext in [".docx", ".doc"]:
            chunks = load_docx(str(file_path))
        elif file_ext in [".txt", ".md"]:
            chunks = load_text_file(str(file_path))
        
        if chunks:
            if collection_name == "vnd":
                from vnd_masking import mask_vnd_chunks

                chunks = mask_vnd_chunks(chunks)
            # Генерируем уникальные ID на основе имени файла и индекса
            import hashlib
            file_hash = hashlib.md5(str(file_path).encode()).hexdigest()[:8]
            metadata = [{"source": str(file_path), "filename": file_path_obj.name} for _ in chunks]
            ids = [f"{collection_name}_{file_hash}_{i}" for i in range(len(chunks))]
            
            print(f"  Добавление {len(chunks)} чанков в векторную базу '{collection_name}'...")
            store.add_documents(chunks, metadata, ids)
            print(f"  ✓ Обработан: {len(chunks)} чанков из файла {file_path_obj.name}")
            print(f"  ✓ Файл добавлен в векторную базу с ID: {collection_name}_{file_hash}_*")
            
            # Проверяем, что документы действительно добавлены
            try:
                collection_info = store.get_collection_info()
                print(f"  ✓ Всего документов в базе '{collection_name}': {collection_info['count']}")
            except:
                pass
            
            return {
                "status": "success",
                "files_processed": 1,
                "total_chunks": len(chunks),
                "files_found": 1,
                "files_skipped": 0,
                "filename": file_path_obj.name,
                "collection": collection_name,
                "message": f"Файл успешно обработан и добавлен в векторную базу. Создано {len(chunks)} фрагментов."
            }
        else:
            error_msg = f"Файл не содержит текста: {file_path_obj.name}"
            print(f"  ⚠ {error_msg}")
            return {
                "status": "warning",
                "message": error_msg,
                "files_processed": 0,
                "total_chunks": 0,
                "files_found": 1,
                "files_skipped": 1
            }
    except Exception as e:
        error_msg = f"Ошибка при обработке {file_path_obj.name}: {e}"
        print(f"  ✗ {error_msg}")
        import traceback
        traceback.print_exc()
        return {
            "status": "error",
            "message": error_msg,
            "files_processed": 0,
            "total_chunks": 0,
            "files_found": 1,
            "files_skipped": 1
        }

def process_folder(folder_path: str, store, collection_name: str) -> Dict:
    """Обработка папки с документами"""
    folder_path_obj = Path(folder_path)
    
    if not folder_path_obj.exists():
        error_msg = f"Папка {folder_path} не существует"
        print(error_msg)
        return {"status": "error", "message": error_msg}
    
    if not folder_path_obj.is_dir():
        error_msg = f"{folder_path} не является папкой"
        print(error_msg)
        return {"status": "error", "message": error_msg}
    
    print(f"Обработка папки: {folder_path}")
    files_processed = 0
    total_chunks = 0
    files_found = 0
    files_skipped = []
    
    # Поддерживаемые расширения
    supported_extensions = {".pdf", ".docx", ".doc", ".txt", ".md"}
    
    for file_path in folder_path_obj.rglob("*"):
        if file_path.is_file():
            files_found += 1
            file_ext = file_path.suffix.lower()
            
            # Пропускаем файлы без расширения или с неподдерживаемым расширением
            if not file_ext or file_ext not in supported_extensions:
                files_skipped.append(str(file_path))
                continue
            
            print(f"Обработка файла: {file_path.name}")
            chunks = []
            
            try:
                if file_ext == ".pdf":
                    chunks = load_pdf(str(file_path))
                elif file_ext in [".docx", ".doc"]:
                    chunks = load_docx(str(file_path))
                elif file_ext in [".txt", ".md"]:
                    chunks = load_text_file(str(file_path))
                
                if chunks:
                    if collection_name == "vnd":
                        from vnd_masking import mask_vnd_chunks

                        chunks = mask_vnd_chunks(chunks)
                    import hashlib
                    file_hash = hashlib.md5(str(file_path).encode()).hexdigest()[:8]
                    metadata = [{"source": str(file_path), "filename": file_path.name} for _ in chunks]
                    ids = [f"{collection_name}_{file_hash}_{i}" for i in range(len(chunks))]
                    store.add_documents(chunks, metadata, ids)
                    files_processed += 1
                    total_chunks += len(chunks)
                    print(f"  ✓ Обработан: {len(chunks)} чанков")
                else:
                    print(f"  ⚠ Файл не содержит текста: {file_path.name}")
                    files_skipped.append(str(file_path))
            except Exception as e:
                print(f"  ✗ Ошибка при обработке {file_path.name}: {e}")
                files_skipped.append(str(file_path))
    
    result = {
        "status": "success",
        "files_processed": files_processed,
        "total_chunks": total_chunks,
        "files_found": files_found,
        "files_skipped": len(files_skipped)
    }
    
    if files_skipped:
        result["skipped_files"] = files_skipped[:10]  # Первые 10 для примера
    
    print(f"Итоги обработки папки {folder_path}: обработано {files_processed} файлов, {total_chunks} чанков")
    
    return result

def create_index_bases():
    """Создание индексных баз из документов"""
    from config import settings
    import vector_store
    
    # Создаем папки, если их нет
    os.makedirs(settings.fz_folder, exist_ok=True)
    os.makedirs(settings.fzyur_folder, exist_ok=True)
    os.makedirs(settings.in_folder, exist_ok=True)
    os.makedirs(settings.out_folder, exist_ok=True)
    
    # Инициализируем векторные базы
    vector_store.init_vector_stores()
    
    results = {}
    
    # ПРИОРИТЕТ 1: Загружаем ФЗ из FZ (основной источник)
    print("=" * 60)
    print("Индексация базы ФЗ (основной источник)...")
    print("=" * 60)
    if vector_store.fz_store:
        results["fz"] = process_folder(settings.fz_folder, vector_store.fz_store, "fz")
        if results["fz"].get("status") == "success" and results["fz"].get("files_processed", 0) > 0:
            print(f"✓ База ФЗ успешно загружена: {results['fz'].get('files_processed')} файлов")
        else:
            print(f"⚠ База ФЗ: {results['fz'].get('message', 'файлы не найдены или не обработаны')}")
    else:
        results["fz"] = {"status": "error", "message": "База ФЗ не инициализирована"}
        print("✗ Ошибка: База ФЗ не инициализирована")
    
    # ПРИОРИТЕТ 2: Загружаем ГОСТ из FZYur (опционально)
    print("\n" + "=" * 60)
    print("Индексация базы ГОСТ (опционально)...")
    print("=" * 60)
    if vector_store.gost_store:
        results["gost"] = process_folder(settings.fzyur_folder, vector_store.gost_store, "gost")
        if results["gost"].get("status") == "success" and results["gost"].get("files_processed", 0) > 0:
            print(f"✓ База ГОСТ успешно загружена: {results['gost'].get('files_processed')} файлов")
        else:
            print(f"ℹ База ГОСТ: {results['gost'].get('message', 'файлы не найдены или не обработаны')} (не критично)")
    else:
        results["gost"] = {"status": "warning", "message": "База ГОСТ не инициализирована (не критично)"}
        print("ℹ База ГОСТ не инициализирована (не критично для анализа)")
    
    # Загружаем ВНД из IN (для анализа)
    print("\n" + "=" * 60)
    print("Индексация базы ВНД...")
    print("=" * 60)
    if vector_store.vnd_store:
        results["vnd"] = process_folder(settings.in_folder, vector_store.vnd_store, "vnd")
        if results["vnd"].get("status") == "success" and results["vnd"].get("files_processed", 0) > 0:
            print(f"✓ База ВНД успешно загружена: {results['vnd'].get('files_processed')} файлов")
        else:
            print(f"ℹ База ВНД: {results['vnd'].get('message', 'файлы не найдены')}")
    else:
        results["vnd"] = {"status": "error", "message": "База ВНД не инициализирована"}
        print("✗ Ошибка: База ВНД не инициализирована")
    
    print("\n" + "=" * 60)
    print("Индексация завершена")
    print("=" * 60)
    
    return results

