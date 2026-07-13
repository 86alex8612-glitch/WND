"""
Помощник в создании ВНД: переработка существующего документа и создание нового.
"""
from __future__ import annotations

import json
import io
import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.shared import RGBColor
from openai import OpenAI

from config import settings
from document_loader import extract_full_text

logger = logging.getLogger("create_vnd")

GOST_QUERY = (
    "ГОСТ Р 7.0.97-2025 Система стандартов по информации "
    "организационно-распорядительная документация требования к оформлению"
)

REWORK_PROMPT = """Ты — старший юрисконсульт методист с опытом работы более 10 лет.

ГЛАВНАЯ ЗАДАЧА: подготовить ПОЛНОЦЕННЫЙ переработанный текст локального нормативного акта (ВНД), в котором УСТРАНЕНЫ ВСЕ недостатки из отчёта анализа. Переработанный документ должен СУЩЕСТВЕННО отличаться от исходного там, где анализ выявил пробелы, несоответствия или отсутствующие нормы.

ЗАПРЕЩЕНО:
- возвращать исходный текст с минимальными правками или только новыми заголовками;
- оставлять главы без развёрнутого содержания;
- игнорировать пункты плана переработки и отчёта анализа.

ОБЯЗАТЕЛЬНЫЙ АЛГОРИТМ:
1. Внимательно изучи ОТЧЁТ АНАЛИЗА (загруженный пользователем или сформированный ранее). Это главный источник изменений.
2. Пройди по КАЖДОМУ пункту «ПЛАН ПЕРЕРАБОТКИ» и отчёта анализа — для каждого создай, измени или удали соответствующую статью.
3. Внеси ВСЕ изменения и дополнения из отчёта анализа: каждая рекомендация, каждый недостаток и каждое отсутствующее требование ФЗ должны быть отражены в переработанном ВНД.
4. Если в отчёте указано удалить, объединить или заменить норму — выполни это в тексте документа.
5. Сопоставь исходный ВНД с основным федеральным законом (из контекста ФЗ): для статей ФЗ, не отражённых в ВНД, — добавь новые статьи с адаптированным содержанием.
6. Для каждой выявленной проблемы в отчёте — включи конкретную норму, устраняющую эту проблему.

НОРМАТИВНЫЕ АКТЫ ИЗ ОТЧЁТА АНАЛИЗА (ТИП Б и аналогичные блоки):
- Если в отчёте анализа указан федеральный закон, ГОСТ, подзаконный акт или иной НПА, которого НЕТ в разделе «используемые документы» / «нормативные ссылки» / «перечень НПА» исходного ВНД:
  1) ОБЯЗАТЕЛЬНО добавь этот документ в «Список нормативных правовых актов, использованных при разработке» (полное наименование и реквизиты);
  2) исправь текст ВНД согласно отчёту анализа — добавь или измени статьи, отражающие требования этого НПА.
- Каждый пункт отчёта «не учтены требования …» / «что должно быть в документе» должен быть устранён новой или изменённой статьёй.

Пример устранения недостатка (ТИП Б — неучтённый ГОСТ Р 57580.1-2017):
• В список используемых документов добавить: ГОСТ Р 57580.1-2017 «Безопасность финансовых (банковских) операций. Защита информации…».
• Добавить статью, например «Статья X.Y. Меры защиты информации»:
  [преамбула: в соответствии с требованиями ГОСТ Р 57580.1-2017 организация обеспечивает установленные меры защиты информации…]
  <<RED>>Необходимо указать самостоятельно перечень мер защиты информации, принятых в организации.<<ENDRED>>

ЗАПОЛНЕНИЕ ДАННЫХ, ИЗВЕСТНЫХ ТОЛЬКО ОРГАНИЗАЦИИ:
Если невозможно указать конкретные сведения (мероприятия на конкретном предприятии, локальные процедуры, перечни, ответственные лица, сроки внедрения и т.п.):
- напиши статью с нормативной преамбулой (правовое основание, общие требования из отчёта анализа и НПА);
- затем ОБЯЗАТЕЛЬНО добавь маркер (без изменения текста маркера):
  <<RED>>Необходимо указать самостоятельно [конкретизируй, что именно должен заполнить заказчик].<<ENDRED>>
- не выдумывай фактические данные организации; не оставляй пробел — всегда используй маркер <<RED>>…<<ENDRED>>.

СТРУКТУРА И ОФОРМЛЕНИЕ (по ГОСТ из контекста):
Документ строится иерархически:
  Глава 1. [Название]
    Статья 1.1. [Название статьи]
    [Развёрнутый нормативный текст: минимум 3–5 предложений с конкретными обязанностями, процедурами, сроками, ответственными]
    Статья 1.2. ...

Обязательные главы (если отсутствуют — создай с полным текстом статей):
- Глава «Общие положения» (цели, задачи, правовое основание, ссылка на Трудовой кодекс РФ или иной ФЗ, дающий право издавать документ).
- Раздел «Список нормативных правовых актов, использованных при разработке» (все ФЗ, ГОСТ и иные НПА из исходного ВНД и из отчёта анализа).
- Глава «Термины и определения» (не менее 5 терминов с определениями).
- Глава с порядком основного процесса (пошаговые статьи).
- Глава «Права и обязанности» (отдельные статьи для работодателя/работника или иных субъектов).
- Глава «Ответственность».
- Глава «Заключительные положения» (вступление в силу, порядок изменений).

СОДЕРЖАНИЕ СТАТЕЙ:
- Каждая статья — самостоятельная норма с подлежащим, действием, условиями, сроками.
- Используй нумерацию пунктов: 1), 2), 3) или 1.1, 1.2 внутри статьи при необходимости.
- Ссылайся на статьи ФЗ в формулировках «в соответствии со статьёй N Федерального закона … организация обязана …».
- Требования федерального закона переложи в обязанности организации и порядок действий сотрудников.

СТИЛЬ: официально-деловой, юридически грамотный, без двусмысленностей.

ДОПОЛНИТЕЛЬНО:
- Проверь актуальность ссылок на ФЗ; замени устаревшие.
- Дату утверждения — маскированная текущая дата (**********, по одному * на символ ДД.ММ.ГГГГ).
- Год утверждения/введения в действие — только текущий календарный год (не копируй год из исходного документа).
- ФИО, наименование организации-работодателя (ООО «…»), адреса, ИНН, ОГРН и иные реквизиты организации — символ * (по одному * на каждый символ).
- НЕ маскируй ссылки на нормативные акты: указывай полностью Федеральные законы, ГОСТ, положения и указания Банка России, постановления Правительства РФ (тип, дата, номер, название в «кавычках»).

ФОРМАТ ОТВЕТА: только полный текст переработанного ВНД от названия до заключительных положений. Без комментариев «что изменено». Маркеры <<RED>>…<<ENDRED>> включай в текст статьи как указано выше."""

RED_MARKER_PATTERN = re.compile(r"<<RED>>(.*?)<<ENDRED>>", re.DOTALL)

REWORK_CHANGES_REPORT_PROMPT = """Ты — юрисконсульт. Составь отчёт о переработке ВНД, сравнив исходный документ, отчёт анализа и переработанный текст.

Используй ТОЛЬКО следующий шаблон (без других разделов):

ОТЧЁТ О ПЕРЕРАБОТКЕ ВНД
Документ: {document_name}
Дата: {report_date}

Удалено — статьи:
• [номер и название статьи] — [краткое обоснование; ссылка на отчёт анализа]
(если удалений нет — одна строка: нет)

Изменены — статьи:
• [номер и название статьи] — [что изменено; ссылка на пункт отчёта анализа]
(если изменений нет — одна строка: нет)

Добавлены — статьи:
• [номер и название статьи] — [основание: рекомендация из отчёта анализа / статья ФЗ]
(если добавлений нет — одна строка: нет)

Указывай конкретные номера статей (Статья 1.1, Статья 2.3 и т.д.). Каждое изменение и дополнение связывай с отчётом анализа."""

NEW_VND_PROMPT = """Ты — эксперт по организационно-распорядительной документации.

Подготовь проект нового внутреннего нормативного документа (ВНД) организации.

Учитывай:
- сферу деятельности и форму собственности организации;
- области законодательства, указанные заказчиком;
- применимые федеральные законы из контекста;
- требования ГОСТ Р 7.0.97-2025 к оформлению ОРД;
- дату создания/утверждения документа указывай маскированной текущей датой (********** — по одному * на символ формата ДД.ММ.ГГГГ).
- год утверждения/введения в действие указывай текущим календарным годом.

Документ должен быть структурированным, юридически корректным и готовым к согласованию.
Формат ответа: полный текст документа (без пояснений вне документа)."""

NEW_VND_ANALYSIS_PROMPT = """Ты — эксперт-юрист в области корпоративного права и комплаенса, специализирующийся на создании внутренних нормативных документов для российских организаций.

ЗАДАЧА: Проанализировать предоставленные вводные данные о планируемом к созданию документе.

ИНСТРУКЦИЯ ПО АНАЛИЗУ:
- Проанализируй каждую из переменных.
- На основе анализа определи, какие федеральные законы, кодексы, постановления Правительства РФ и национальные стандарты (ГОСТ) являются основополагающими для регулирования указанной сферы.
- Учитывай специфику сферы деятельности (например, для банковской сферы — ФЗ «О банках и банковской деятельности», для образования — ФЗ «Об образовании в РФ») и формы собственности (например, ФЗ «Об обществах с ограниченной ответственностью» для ООО).
- Если отношение к гостайне указано как «Да», обязательно включи в список ФЗ «О государственной тайне».

ФОРМАТ ВЫВОДА:
1. Краткое резюме анализа: для каких целей и на основе каких вводных будет создаваться документ.
2. Пронумерованный список только названий федеральных нормативных правовых актов и стандартов (без ссылок), которые необходимо использовать в качестве правовой базы. Список должен быть исчерпывающим для данной задачи."""

NEW_VND_GENERATION_PROMPT = """Ты — высококвалифицированный юрист-методолог. Твоя задача — составить юридически грамотный, структурированный и стилистически выверенный проект внутреннего нормативного документа.

ЗАДАЧА: Создать проект документа на основе предоставленных материалов.

ИНСТРУКЦИЯ ПО СОЗДАНИЮ:
- Внимательно изучи Текст_документов и ГОСТ_Р_7_0_97_2025.
- Строго соблюдай требования ГОСТ Р 7.0.97-2025 к оформлению организационно-распорядительной документации (структура, реквизиты, заголовок, нумерация).
- Структурируй документ по схеме:
  • Общие положения (цели, задачи, нормативная база, область применения);
  • Основные термины и определения;
  • Основная часть (процессы, права, обязанности, процедуры — по области законодательства);
  • Ответственность;
  • Порядок ознакомления;
  • Приложения (при необходимости).
- Используй деловой стиль русского языка. Формулировки чёткие и исполнимые.

ОРГАНИЗАЦИЯ:
- Во всём документе (шапка, преамбула, область применения, права и обязанности) указывай организацию ТОЛЬКО под наименованием из вводных данных («Наименование организации в документе»), например: ООО «DialogAI».
- Не заменяй название организации звёздочками и не используй обобщения «Организация ***».

ТЕРМИНОЛОГИЯ (БЕЗ ИЗБЫТОЧНОГО МАСКИРОВАНИЯ):
- Пиши полные юридические термины и названия процессов: «обработка персональных данных», «хранение персональных данных», «коммерческая тайна», «информационная безопасность», «трудовой договор» и т.п.
- ЗАПРЕЩЕНО маскировать звёздочками (*) понятия, процедуры, виды информации, названия разделов и статей.
- Пример НЕВЕРНО: «2.2. Обработка ***************». Пример ВЕРНО: «2.2. Обработка персональных данных».

ЧТО МОЖНО НЕ УКАЗЫВАТЬ / ВЫДЕЛЯТЬ:
- Конкретные ФИО должностных лиц — указывай должность («Генеральный директор», «Ответственный за обработку ПДн») без вымышленных ФИО.
- Адреса, ИНН, ОГРН, КПП, банковские реквизиты, телефоны, e-mail — не выдумывай; для таких сведений используй маркер: <<RED>>Необходимо указать самостоятельно [что именно].<<ENDRED>>
- Дату утверждения указывай маскированной текущей датой (********** — по одному * на символ ДД.ММ.ГГГГ).
- Год утверждения/введения в действие — текущий календарный год.

ОБЪЁМ И ПОЛНОТА:
- Документ должен быть ПОЛНЫМ локальным актом, а не кратким шаблоном.
- Минимум 15–25 статей/пунктов в основной части суммарно.
- Каждый пункт обязательного чек-листа должен быть раскрыт отдельным подразделом или статьёй.
- Используй уточняющие ответы пользователя как обязательные факты (не выдумывай противоречащее).

ФОРМАТ ВЫВОДА: полный текст проекта внутреннего нормативного документа, готовый к использованию."""

NEW_VND_PLAN_PROMPT = """Ты — юрист-методолог. Составь детальный план (оглавление) внутреннего нормативного документа.

ЗАДАЧА: На основе вводных данных, уточняющих ответов пользователя, результата правового анализа и обязательного чек-листа
сформировать структуру документа с разделами для поэтапной генерации.

ТРЕБОВАНИЯ:
- Каждый пункт чек-листа должен быть отражён в одном из разделов.
- Учитывай ответы пользователя (трансграничная передача, категории субъектов и т.д.).
- Для каждого раздела укажи: id (латиница), title (заголовок с нумерацией), requirements (что включить), min_articles (минимум статей/пунктов).

ФОРМАТ: только JSON без markdown:
{
  "sections": [
    {"id": "general", "title": "1. Общие положения", "requirements": "...", "min_articles": 4}
  ]
}"""

NEW_VND_SECTION_PROMPT = """Ты — юрист-методолог. Пиши ТОЛЬКО один раздел внутреннего нормативного документа.

ТРЕБОВАНИЯ:
- Строго соблюдай ГОСТ Р 7.0.97-2025 (нумерация, стиль).
- Используй наименование организации из вводных.
- Пиши полные юридические термины без маскирования звёздочками.
- Каждая статья/пункт — развёрнутый текст (минимум 3–5 предложений в статье).
- Учитывай уточняющие ответы пользователя и чек-лист.
- Не повторяй другие разделы — только текущий.
- Для сведений, известных только организации: <<RED>>Необходимо указать самостоятельно [что именно].<<ENDRED>>

ФОРМАТ: только текст раздела с заголовком, без комментариев."""

NEW_VND_RECOMMENDATIONS_PROMPT = """Ты — юрист-методолог. Составь завершающий раздел «Рекомендации» для проекта внутреннего нормативного документа.

ЗАДАЧА: Перечислить дополнительные внутренние нормативные документы, которые организации целесообразно разработать в дополнение к уже созданному документу.

ТРЕБОВАНИЯ:
- Раздел носит рекомендательный характер (не является обязательной частью утверждаемого текста, но включается в проект для заказчика).
- Учитывай область законодательства, сферу деятельности, уточняющие ответы пользователя и содержание уже сгенерированного документа.
- Не дублируй сам создаваемый документ — перечисляй СОПУТСТВУЮЩИЕ акты (положения, приказы, регламенты, инструкции, формы, соглашения, модели угроз и т.п.).
- Минимум 6 позиций, максимум 12.
- Для каждой позиции укажи: вид документа, приблизительное наименование, степень конфиденциальности, краткое назначение (1 предложение).
- Степень конфиденциальности — одна из: Общедоступный; Конфиденциально (внутренний документ); Для служебного пользования (ДСП); Коммерческая тайна; Персональные данные (ограниченный доступ); Сведения, составляющие государственную тайну.
- Если отношение к гостайне — «Да», включи хотя бы один документ с соответствующим уровнем защиты.
- Пиши полные термины без маскирования звёздочками.

ФОРМАТ ВЫВОДА:
Заголовок: «Рекомендации по разработке дополнительных документов»
Краткое вступление (2–3 предложения).
Далее нумерованный перечень в виде:
N. [Вид документа] «[Приблизительное наименование]»
   Степень конфиденциальности: ...
   Назначение: ...

Без таблиц markdown, без комментариев вне раздела."""

CREATE_QA_MODE_SUFFIX = """
РЕЖИМ ДИАЛОГА ПО ПОДГОТОВЛЕННОМУ ВНД:
Ты продолжаешь работу в той же роли и с теми же компетенциями, что при создании документа.
Сейчас отвечаешь на вопросы пользователя ТОЛЬКО по подготовленному ВНД (новому документу), текст которого приведён ниже.

ПРАВИЛА ОТВЕТОВ:
- Отвечай по существу, ссылайся на конкретные статьи и главы документа.
- Поясняй нормы, правовые основания, порядок применения, места с формулировкой «Необходимо указать самостоятельно».
- Не переписывай весь документ заново — отвечай на конкретный вопрос.
- Если вопрос НЕ относится к этому ВНД (другие документы, общие темы не по документу, посторонние запросы) — вежливо откажи и предложи задать вопрос по подготовленному документу.
- Не выдумывай положения, отсутствующие в тексте; при необходимости укажи, что сведения подлежат самостоятельному заполнению организацией.

ФОРМАТ ОТВЕТА: краткий развёрнутый ответ консультанта (без повторения всего документа)."""

CREATE_QA_WELCOME = (
    "Здравствуйте! Задайте вопрос по подготовленному документу — "
    "поясню статьи, структуру, правовые основания и места для самостоятельного заполнения."
)


def get_in_folder() -> Path:
    """Папка IN — загруженные ВНД (config.cfg: data_root / data_win)."""
    return Path(settings.in_folder)


def get_out_folder() -> Path:
    """Папка OUT — отчёты анализа."""
    return Path(settings.out_folder)


def _resolve_vnd_file_path(filename: str) -> Path:
    """Найти ВНД только в IN/."""
    safe = os.path.basename((filename or "").strip())
    if not safe:
        raise FileNotFoundError("Файл не указан")

    in_path = get_in_folder() / safe
    if in_path.is_file():
        return in_path

    original = re.sub(r"^main_\d{8}_\d{6}_", "", safe, flags=re.IGNORECASE)
    if original and original != safe:
        alt = get_in_folder() / original
        if alt.is_file():
            return alt

    from federal_refs import find_vnd_file

    try:
        return find_vnd_file(safe)
    except FileNotFoundError as exc:
        raise FileNotFoundError(
            f"ВНД не найден в {get_in_folder()}: {safe}. "
            "Загрузите документ вручную или нажмите «Взять из IN (после анализа)»."
        ) from exc


def _save_to_in_folder(filename: str, content: bytes) -> Path:
    """Сохранить ВНД в IN/."""
    safe = os.path.basename(filename or "document")
    in_folder = get_in_folder()
    in_folder.mkdir(parents=True, exist_ok=True)
    target = in_folder / safe
    target.write_bytes(content)
    return target


def _save_to_out_folder(filename: str, content: bytes) -> Path:
    """Сохранить отчёт анализа в OUT/."""
    safe = os.path.basename(filename or "analysis")
    out_folder = get_out_folder()
    out_folder.mkdir(parents=True, exist_ok=True)
    target = out_folder / safe
    target.write_bytes(content)
    return target


def import_main_from_in(filename: str) -> dict:
    """Подключить ВНД из IN/ (без копирования в другие папки)."""
    in_path = _resolve_vnd_file_path(filename)
    return {
        "filename": in_path.name,
        "file_path": str(in_path),
        "in_folder": str(get_in_folder()),
        "kind": "main",
    }


def save_create_upload(filename: str, content: bytes, kind: str) -> dict:
    """Загрузить файл: main → IN/, analysis → OUT/."""
    if kind == "main":
        return save_main_for_rework(filename, content)

    out_path = _save_to_out_folder(filename, content)
    return {
        "filename": out_path.name,
        "file_path": str(out_path),
        "out_folder": str(get_out_folder()),
        "kind": "analysis",
        "source": "out",
    }


def save_main_for_rework(filename: str, content: bytes) -> dict:
    """Сохранить основной ВНД для переработки в IN/."""
    safe = os.path.basename(filename or "document")
    in_path = _save_to_in_folder(safe, content)
    return {
        "filename": in_path.name,
        "file_path": str(in_path),
        "in_folder": str(get_in_folder()),
        "kind": "main",
    }


def read_create_file(filename: str, max_chars: int = 80000) -> str:
    """Прочитать текст ВНД из IN/."""
    path = _resolve_vnd_file_path(filename)
    return extract_full_text(str(path))[:max_chars]


def read_analysis_report_file(filename: str, max_chars: int = 50000) -> str:
    """Прочитать текст отчёта анализа из OUT/."""
    from search_vnd import resolve_analysis_text

    text, _, _ = resolve_analysis_text(filename, "out")
    return text[:max_chars]


def get_create_options() -> dict:
    from new_vnd_form import get_new_vnd_form_options
    from pre_analysis import ACTIVITY_SPHERES, OWNERSHIP_FORMS

    return {
        "explanation": (
            "ВНД — это внутренний (локальный) нормативный документ организации."
        ),
        "activity_spheres": ACTIVITY_SPHERES,
        "ownership_forms": OWNERSHIP_FORMS,
        "new_document": get_new_vnd_form_options(),
    }


def _get_openai_client() -> OpenAI:
    if not settings.openai_api_key:
        raise RuntimeError("OPENAI_API_KEY не установлен")
    return OpenAI(api_key=settings.openai_api_key)


def get_gost_context() -> str:
    from analiz import get_relevant_context
    from vector_store import gost_store, init_vector_stores

    parts = []
    try:
        init_vector_stores()
        if gost_store:
            results = gost_store.search(GOST_QUERY, n_results=4)
            for item in results or []:
                name = item["metadata"].get("filename", "ГОСТ")
                if "7.0.97" in name or "7.0.97" in (item.get("document") or ""):
                    parts.append(f"{name}\n{item['document'][:1200]}")
    except Exception as exc:
        logger.warning("Поиск ГОСТ в базе: %s", exc)

    rag = get_relevant_context(GOST_QUERY, max_results=3, prioritize_vnd=False)
    if rag and "не найден" not in rag.lower():
        parts.append(rag[:2500])

    return "\n\n".join(parts) if parts else "ГОСТ Р 7.0.97-2025 — требования к оформлению ОРД."


def get_fz_context(query: str, max_results: int = 4) -> str:
    from vector_store import fz_store, init_vector_stores

    try:
        init_vector_stores()
        if not fz_store:
            return ""
        results = fz_store.search(query, n_results=max_results)
        chunks = []
        for item in results or []:
            name = item["metadata"].get("filename", "ФЗ")
            chunks.append(f"{name}\n{item['document'][:800]}")
        return "\n\n".join(chunks)
    except Exception as exc:
        logger.warning("Поиск ФЗ: %s", exc)
        return ""


def _llm_json(prompt: str, max_tokens: int = 800) -> dict:
    client = _get_openai_client()
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Отвечай только валидным JSON без markdown."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
        max_tokens=max_tokens,
    )
    raw = (response.choices[0].message.content or "").strip()
    raw = re.sub(r"^```json\s*|\s*```$", "", raw, flags=re.IGNORECASE)
    return json.loads(raw)


def identify_required_federal_laws(stage1: dict) -> List[dict]:
    """Определить необходимые ФЗ по данным шага 1 (вопрос 3 — свободный ввод)."""
    legal_text = (
        stage1.get("legal_areas_text")
        or ", ".join(stage1.get("legal_areas") or [])
        or stage1.get("document_purpose")
        or ""
    )
    prompt = f"""По параметрам организации определи федеральные законы и кодексы РФ,
необходимые для подготовки внутреннего нормативного документа.

Сфера деятельности: {stage1.get('activity_sphere', '')}
Форма собственности: {stage1.get('ownership_form', '')}
Тема/области законодательства (запрос заказчика): {legal_text}

Верни JSON:
{{
  "laws": [
    {{
      "title": "краткое название",
      "number": "152-ФЗ или пусто",
      "search_query": "строка для поиска на pravo.gov.ru"
    }}
  ]
}}
От 2 до 8 наиболее релевантных документов."""

    try:
        data = _llm_json(prompt, max_tokens=900)
        laws = data.get("laws") or []
        normalized = []
        for law in laws:
            if not isinstance(law, dict):
                continue
            normalized.append({
                "title": law.get("title") or law.get("search_query") or "Федеральный закон",
                "number": law.get("number") or "",
                "search_query": law.get("search_query") or law.get("title") or "",
                "raw": law.get("search_query") or law.get("title") or "",
            })
        return normalized
    except Exception as exc:
        logger.error("Ошибка определения ФЗ: %s", exc)
        return [{
            "title": legal_text[:120] or "Федеральный закон",
            "number": "",
            "search_query": legal_text[:200] or "федеральный закон",
            "raw": legal_text[:200],
        }]


def download_required_laws(laws: List[dict]) -> dict:
    from pravo_downloader import download_and_index_references

    refs = []
    for law in laws:
        refs.append({
            "title": law.get("title") or "Федеральный закон",
            "number": law.get("number") or "",
            "search_query": law.get("search_query") or law.get("title") or "",
            "raw": law.get("raw") or law.get("search_query") or "",
        })
    return download_and_index_references(refs)


def _prepare_federal_refs_context(main_filename: str, main_text: str) -> tuple[str, dict]:
    """Контекст по ссылкам в документе + попытка подкачки отсутствующих."""
    from federal_refs import detect_federal_references_from_file

    path = _resolve_vnd_file_path(main_filename)
    refs_result = detect_federal_references_from_file(str(path))
    missing = refs_result.get("missing_references") or []

    download_info = {"status": "skipped", "message": "Все ссылки есть в локальной базе"}
    if missing:
        try:
            download_info = download_required_laws(missing)
        except Exception as exc:
            logger.warning("Не удалось подкачать ФЗ: %s", exc)
            download_info = {"status": "error", "message": str(exc)}

    lines = ["=== ССЫЛКИ НА ФЕДЕРАЛЬНЫЕ ДОКУМЕНТЫ В ИСХОДНОМ ВНД ==="]
    for ref in refs_result.get("unique_references") or refs_result.get("references") or []:
        status = "есть в базе" if ref.get("in_local_base") else "отсутствовал в базе"
        lines.append(
            f"- {ref.get('title', ref.get('raw', ''))} ({ref.get('number', '')}) — {status}"
        )
    lines.append(f"\n{download_info.get('message', '')}")
    return "\n".join(lines), {"refs": refs_result, "download": download_info}


def _ensure_rework_federal_refs(main_filename: str) -> dict:
    """Подготовить федеральные ссылки как в карточке «Анализ ВНД» (без сохранения отчёта)."""
    path = _resolve_vnd_file_path(main_filename)
    from federal_refs import detect_federal_references_from_file

    refs_result = detect_federal_references_from_file(str(path))
    missing = refs_result.get("missing_references") or []
    download_info = {"status": "skipped", "message": "Все ссылки есть в локальной базе"}
    if missing:
        try:
            download_info = download_required_laws(missing)
        except Exception as exc:
            logger.warning("Не удалось подкачать ФЗ перед анализом: %s", exc)
            download_info = {"status": "error", "message": str(exc)}
    return {"refs": refs_result, "download": download_info}


def run_document_analysis(
    main_filename: str,
    vnd_name: str,
    stage1: dict,
) -> tuple[str, dict, dict]:
    """Провести анализ как в блоке «Анализ ВНД» (отчёт не сохраняется на диск)."""
    from analiz import analyze_vnd
    from pre_analysis import normalize_stage1_answers

    refs_meta = _ensure_rework_federal_refs(main_filename)

    main_text = read_create_file(main_filename)
    pre_analysis = normalize_stage1_answers(stage1)
    result = analyze_vnd(
        user_message="Приступи к правовому анализу загруженного ВНД.",
        history=[],
        vnd_text=main_text,
        force_analysis=True,
        pre_analysis=pre_analysis,
    )
    if isinstance(result, dict):
        content = result.get("content", "")
    else:
        content = str(result)
    return content, pre_analysis, refs_meta


def analyze_for_rework(
    main_filename: str,
    vnd_name: str = "",
    stage1: Optional[dict] = None,
) -> dict:
    """Правовой анализ для переработки: как «Анализ ВНД», без сохранения отчёта в файл."""
    if not stage1:
        raise ValueError("Для анализа без отчёта необходимы параметры этапа 1")

    analysis_text, normalized, refs_meta = run_document_analysis(
        main_filename, vnd_name, stage1
    )

    return {
        "analysis_text": analysis_text,
        "stage1": normalized,
        "analysis_meta": {
            "stage1": normalized,
            "auto_analysis": True,
            "saved_to_file": False,
            "federal_refs": refs_meta,
        },
        "vnd_name": vnd_name,
    }


def detect_rework_stage1(
    main_filename: str,
    vnd_name: str,
    analysis_text: str = "",
    analysis_filename: str = "",
) -> dict:
    from pre_analysis import detect_stage1

    main_text = read_create_file(main_filename)
    safe_name = os.path.basename(main_filename)
    combined = main_text
    extra = (analysis_text or "").strip()
    if not extra and analysis_filename:
        try:
            extra = read_analysis_report_file(analysis_filename)
        except Exception:
            extra = ""
    if extra:
        combined = f"{main_text}\n\n--- Отчёт анализа ---\n{extra[:12000]}"
    return detect_stage1(safe_name, vnd_name or safe_name, combined)


_OLD_SPHERE_TO_FORM_ID = {
    "Производство": "manufacturing",
    "Торговля": "trade_services",
    "Услуги": "trade_services",
    "Финансы": "finance",
    "Строительство": "construction_realestate",
    "Сельское хозяйство": "manufacturing",
    "Информационные технологии (IT)": "it_telecom",
    "Посредническая деятельность": "trade_services",
    "Образовательные услуги": "social",
    "Медицинские услуги": "social",
    "Другое": "trade_services",
}

_FORM_ID_TO_OLD_SPHERE = {
    "finance": "Финансы",
    "it_telecom": "Информационные технологии (IT)",
    "trade_services": "Услуги",
    "manufacturing": "Производство",
    "construction_realestate": "Строительство",
    "social": "Образовательные услуги",
}

_OLD_LEGAL_TO_FORM_ID = {
    "Персональные данные": "personal_data",
    "Информационная безопасность": "confidentiality_ib",
    "Безопасность финансовых (банковских) операций": "finance_risks",
    "Трудовое законодательство": "labor",
    "Противодействие коррупции": "compliance_ethics",
    "Коммерческая тайна": "confidentiality_ib",
    "Государственная тайна": "confidentiality_ib",
    "Образовательная деятельность": "social",
    "Медицинская деятельность": "social",
    "Закупки и контрактная система": "contracts_procurement",
    "Другое": "custom",
}

_FORM_ID_TO_OLD_LEGAL = {
    "corporate": "Другое",
    "personal_data": "Персональные данные",
    "confidentiality_ib": "Информационная безопасность",
    "labor": "Трудовое законодательство",
    "contracts_procurement": "Закупки и контрактная система",
    "finance_risks": "Безопасность финансовых (банковских) операций",
    "compliance_ethics": "Противодействие коррупции",
    "custom": "Другое",
}


def _read_rework_combined_text(
    main_filename: str,
    vnd_name: str = "",
    analysis_text: str = "",
    analysis_filename: str = "",
) -> tuple[str, str]:
    main_text = read_create_file(main_filename)
    safe_name = os.path.basename(main_filename)
    combined = main_text
    extra = (analysis_text or "").strip()
    if not extra and analysis_filename:
        try:
            extra = read_analysis_report_file(analysis_filename)
        except Exception:
            extra = ""
    if extra:
        combined = f"{main_text}\n\n--- Отчёт анализа ---\n{extra[:12000]}"
    return combined, vnd_name or safe_name


def _guess_document_topic(text: str, vnd_name: str) -> str:
    import re

    name = (vnd_name or "").strip()
    if name and len(name) > 8:
        return f"Внутренний нормативный документ: {name[:200]}"
    for pattern in (
        r"(?i)(положение|политика|регламент|инструкция|порядок)[^\n]{0,120}",
        r"(?i)назначение документа[:\s\-–]+([^\n]{10,200})",
        r"(?i)тема[:\s\-–]+([^\n]{10,200})",
    ):
        match = re.search(pattern, text[:8000])
        if match:
            value = (match.group(1) if match.lastindex else match.group(0)).strip()
            if len(value) > 10:
                return value[:300]
    snippet = " ".join(text.split())[:220].strip()
    return snippet or "Внутренний нормативный документ организации"


def _guess_employees_count(text: str) -> str:
    import re

    match = re.search(
        r"(?i)(?:численност\w*|сотрудник\w*|работник\w*)[^\d]{0,20}(\d{1,6})",
        text[:12000],
    )
    return match.group(1) if match else ""


def _guess_branches(text: str) -> str:
    import re

    if re.search(r"(?i)филиал", text[:12000]):
        match = re.search(r"(?i)(\d+\s+филиал\w*|филиал\w*[^\n]{0,80})", text[:12000])
        return (match.group(1) if match else "Есть филиалы/представительства")[:200]
    if re.search(r"(?i)представительств", text[:12000]):
        return "Есть представительства"
    return ""


def _guess_state_secret(text: str) -> str:
    import re

    if re.search(r"(?i)гостайн|государственн\w*\s+тайн", text[:12000]):
        return "yes"
    return "no"


def _guess_target_audience(text: str) -> str:
    import re

    lower = text[:12000].lower()
    if "клиент" in lower and ("сотрудник" in lower or "работник" in lower):
        return "employees_clients"
    if "руководител" in lower or "подразделен" in lower:
        return "managers"
    if "клиент" in lower or "контрагент" in lower:
        return "clients"
    return "all_employees"


def _map_legal_areas_to_form_id(legal_areas: list) -> tuple[str, str]:
    for area in legal_areas or []:
        form_id = _OLD_LEGAL_TO_FORM_ID.get(area)
        if form_id and form_id != "custom":
            return form_id, ""
    if legal_areas:
        return "custom", legal_areas[0]
    return "", ""


def _guess_ownership_form_new(text: str, old_ownership: str) -> str:
    import re

    patterns = [
        (r"(?i)\bООО\b", "ООО (Общество с ограниченной ответственностью)"),
        (r"(?i)\bПАО\b", "ПАО (Публичное акционерное общество)"),
        (r"(?i)\bАО\b", "АО (Непубличное) (Акционерное общество)"),
        (r"(?i)\bИП\b", "ИП (Индивидуальный предприниматель)"),
        (r"(?i)федеральн\w*\s+орган", "Федеральные органы государственной власти"),
        (r"(?i)субъект\w*\s+рф|региональн\w*\s+орган", "Органы государственной власти субъектов РФ"),
        (
            r"(?i)местн\w*\s+самоуправлен|муниципальн\w*\s+орган",
            "Органы местного самоуправления (муниципальные органы)",
        ),
        (r"(?i)государственн\w*\s+(?:предприяти|учрежден|организаци)", "Федеральные органы государственной власти"),
    ]
    for pattern, value in patterns:
        if re.search(pattern, text[:8000]):
            return value
    if old_ownership == "Государственные предприятия":
        return "Федеральные органы государственной власти"
    return "ООО (Общество с ограниченной ответственностью)"


def detect_rework_form_prefill(
    main_filename: str,
    vnd_name: str,
    analysis_text: str = "",
    analysis_filename: str = "",
) -> dict:
    """Параметры этапа 1 + предзаполненная анкета (как «Создать новый»)."""
    from new_vnd_form import get_new_vnd_form_options

    combined, doc_name = _read_rework_combined_text(
        main_filename, vnd_name, analysis_text, analysis_filename
    )
    stage1 = detect_rework_stage1(
        main_filename, vnd_name, analysis_text, analysis_filename
    )

    activity_id = _OLD_SPHERE_TO_FORM_ID.get(stage1.get("activity_sphere") or "", "")
    legal_id, legal_custom = _map_legal_areas_to_form_id(stage1.get("legal_areas") or [])

    clean_name = doc_name
    if "." in clean_name:
        clean_name = os.path.splitext(clean_name)[0]

    form = {
        "document_name": clean_name,
        "document_topic": _guess_document_topic(combined, clean_name),
        "legal_area": legal_id,
        "legal_area_custom": legal_custom,
        "activity_sphere": activity_id,
        "ownership_form": _guess_ownership_form_new(
            combined, stage1.get("ownership_form") or ""
        ),
        "state_secret": _guess_state_secret(combined),
        "employees_count": _guess_employees_count(combined),
        "branches": _guess_branches(combined),
        "target_audience": _guess_target_audience(combined),
        "target_audience_custom": "",
    }

    needs_user = list(stage1.get("needs_user_input") or [])
    for field, value in form.items():
        if field.endswith("_custom"):
            continue
        if not value:
            needs_user.append(field)

    return {
        **stage1,
        "form": form,
        "options": get_new_vnd_form_options(),
        "needs_user_input": sorted(set(needs_user)),
    }


def rework_form_to_stage1(form: dict) -> dict:
    """Преобразовать анкету переработки в параметры этапа 1 для анализа."""
    from pre_analysis import normalize_stage1_answers

    legal_area = (form.get("legal_area") or "").strip()
    if legal_area == "custom":
        legal_label = (form.get("legal_area_custom") or "").strip() or "Другое"
    else:
        legal_label = _FORM_ID_TO_OLD_LEGAL.get(legal_area, "Другое")

    activity_id = (form.get("activity_sphere") or "").strip()
    activity_label = _FORM_ID_TO_OLD_SPHERE.get(activity_id, "Другое")

    ownership = (form.get("ownership_form") or "").strip()
    if ownership.startswith("ООО") or ownership.startswith("АО") or ownership.startswith("ПАО") or ownership.startswith("ИП"):
        ownership_old = "Частные компании"
    elif (
        ownership.startswith("Федеральные")
        or ownership.startswith("Органы государственной власти субъектов")
        or ownership.startswith("Органы местного самоуправления")
        or "государствен" in ownership.lower()
        or "муниципаль" in ownership.lower()
    ):
        ownership_old = "Государственные предприятия"
    else:
        ownership_old = "Частные компании"

    return normalize_stage1_answers({
        "activity_sphere": activity_label,
        "ownership_form": ownership_old,
        "legal_areas": [legal_label] if legal_label else ["Другое"],
    })


def _masked_current_date() -> str:
    """Текущая дата в формате ДД.ММ.ГГГГ, каждый символ заменён на *."""
    current = datetime.now().strftime("%d.%m.%Y")
    return "*" * len(current)


def _current_calendar_year() -> str:
    return str(datetime.now().year)


def _is_law_reference_year(line: str, match_start: int) -> bool:
    """Год в контексте ссылки на НПА — не трогаем."""
    window = line[max(0, match_start - 50): match_start + 30].lower()
    if re.search(r"\d{1,2}\.\d{1,2}\.\d{4}", window):
        return True
    if re.search(
        r"(?:федеральн\w*\s+закон|фз[\s\-№]|закон\s+рф|кодекс|постановлен|гост\s*р?|№\s*\d)",
        window,
    ):
        return True
    return False


def _replace_document_creation_dates(text: str) -> str:
    """Заменить даты/годы утверждения документа на текущие или маскированные значения."""
    masked = _masked_current_date()
    current_year = _current_calendar_year()
    months = (
        "января|февраля|марта|апреля|мая|июня|"
        "июля|августа|сентября|октября|ноября|декабря"
    )

    patterns = [
        # Утверждено ... от ДД.ММ.ГГГГ
        (
            r"(?i)((?:утверждено|утвержден|согласовано)"
            r"(?:\s+(?:приказом|распоряжением|положением))?"
            r"[^\n\d]{0,80}?\s*от\s+)\d{1,2}\.\d{1,2}\.\d{4}"
        ),
        # Дата создания / утверждения / принятия
        (
            r"(?i)((?:дата)\s*(?:создания|утверждения|принятия|разработки|"
            r"введения\s+в\s+действие)?\s*[:\-–—]?\s*)\d{1,2}\.\d{1,2}\.\d{4}"
        ),
        # Приказ/Распоряжение/Положение от ДД.ММ.ГГГГ (шапка документа)
        (
            r"(?i)((?:приказ|распоряжение|положение|инструкция|регламент)"
            r"[^\n]{0,30}?\s*от\s+)\d{1,2}\.\d{1,2}\.\d{4}"
        ),
        # «ДД» месяца ГГГГ г. — типовая форма утверждения
        (
            r"(?i)([«""]?\d{1,2}[»""]?\s+(?:"
            + months
            + r")\s+)\d{4}(\s*г\.?)"
        ),
        # от «ДД» месяца ГГГГ года
        (
            r"(?i)(от\s+[«""]?\d{1,2}[»""]?\s+(?:"
            + months
            + r")\s+)\d{4}(\s*(?:года|г)\.?)"
        ),
    ]

    result = text
    for pattern in patterns:
        result = re.sub(pattern, r"\1" + masked, result)

    # «___»____________2019г. и аналогичные шаблоны с подчёркиваниями/звёздочками
    result = re.sub(
        r"(?i)([«""][^»""\n]*[_]+[^»""\n]*[»""][^\d\n]{0,60})((?:19|20)\d{2})(\s*г\.?)",
        lambda m: m.group(1) + current_year + m.group(3),
        result,
    )
    result = re.sub(
        r"(?i)([_*\s]{4,})((?:19|20)\d{2})(\s*г\.?)",
        lambda m: m.group(1) + current_year + m.group(3),
        result,
    )
    result = re.sub(
        r"(?i)(\*{3,}\s*,\s*)((?:19|20)\d{2})(\s*г\.?)",
        lambda m: m.group(1) + current_year + m.group(3),
        result,
    )

    meta_keywords = re.compile(
        r"(?i)(протокол|утвержден|правлен|редакци|введен|в\s+действие|политик|утверждающ)"
    )
    year_pattern = re.compile(r"(?i)((?:19|20)\d{2})(\s*г\.?)")

    lines = result.split("\n")
    updated_lines = []
    for index, line in enumerate(lines):
        in_header_zone = index < 100
        has_placeholders = bool(re.search(r"[*_]{3,}|«[_\s]+»", line))
        if not (meta_keywords.search(line) or (in_header_zone and has_placeholders)):
            updated_lines.append(line)
            continue

        def _replace_meta_year(match: re.Match) -> str:
            year = match.group(1)
            if year == current_year:
                return match.group(0)
            if _is_law_reference_year(line, match.start()):
                return match.group(0)
            return current_year + match.group(2)

        updated_lines.append(year_pattern.sub(_replace_meta_year, line))

    return "\n".join(updated_lines)


def _mask_fragment(text: str, fragment: str) -> str:
    """Заменить фрагмент звёздочками по числу символов (букв, цифр, пробелов и знаков)."""
    if not fragment or len(fragment) < 2:
        return text
    replacement = "*" * len(fragment)
    return text.replace(fragment, replacement)


def _extract_mask_fragments(document_text: str) -> List[str]:
    """Найти в тексте фрагменты персональных и организационных данных."""
    excerpt = document_text[:12000]
    prompt = f"""Проанализируй организационно-распорядительный документ и найди фрагменты для маскирования:
- ФИО сотрудников и руководителей;
- названия организации-работодателя и её подразделений;
- адреса организации;
- ИНН, ОГРН, КПП, расчётные счета, БИК;
- телефоны, e-mail.

НЕ включай в список ссылки на нормативные акты: Федеральные законы, ГОСТ, положения и указания Банка России, постановления Правительства РФ.

Текст документа:
{excerpt}

Верни JSON:
{{"fragments": ["дословная подстрока из текста", ...]}}

Включай только фрагменты, которые точно встречаются в тексте. Без пояснений."""

    try:
        data = _llm_json(prompt, max_tokens=1200)
        raw = data.get("fragments") or []
        return [str(item).strip() for item in raw if str(item).strip()]
    except Exception as exc:
        logger.warning("Извлечение фрагментов для маскирования: %s", exc)
        return []


def _regex_length_masks(text: str) -> str:
    """Дополнительно замаскировать типовые реквизиты с сохранением длины."""

    def _stars(match: re.Match) -> str:
        return "*" * len(match.group(0))

    patterns = [
        r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}",
        r"(?<!\d)(?:\+7|8)[\s\-]?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}(?!\d)",
        r"ИНН[\s№:]*\d[\d\s]{9,11}",
        r"ОГРН[\s№:]*\d[\d\s]{12,14}",
        r"КПП[\s№:]*\d[\d\s]{8,10}",
        r"БИК[\s№:]*\d[\d\s]{8,10}",
        r"р/с[\s№:]*\d[\d\s]{19,21}",
        r"к/с[\s№:]*\d[\d\s]{19,21}",
    ]
    result = text
    for pattern in patterns:
        result = re.sub(pattern, _stars, result, flags=re.IGNORECASE)
    return result


def mask_personal_placeholders(document_text: str, *, aggressive: bool = True) -> str:
    """Заменить персональные данные. При aggressive=False — только реквизиты и даты (для нового ВНД)."""
    from vnd_masking import (
        is_legal_reference_fragment,
        protect_legal_references,
        restore_legal_references,
    )

    if not document_text.strip():
        return document_text

    protected, placeholders = protect_legal_references(document_text)
    result = protected
    if aggressive:
        try:
            fragments = _extract_mask_fragments(result)
            for fragment in sorted(
                {item for item in fragments if not is_legal_reference_fragment(item)},
                key=len,
                reverse=True,
            ):
                result = _mask_fragment(result, fragment)
        except Exception as exc:
            logger.warning("Маскирование фрагментов: %s", exc)

    result = _regex_length_masks(result)
    result = _replace_document_creation_dates(result)
    return restore_legal_references(result, placeholders)


def _legal_context_hint(form: dict) -> str:
    """Контекст для восстановления терминов вместо звёздочек."""
    return " ".join(
        str(form.get(key) or "")
        for key in (
            "document_name",
            "document_topic",
            "legal_area_resolved",
            "document_name",
        )
    ).lower()


def _term_after_incomplete_heading(word: str, hint: str) -> str:
    """Подобрать продолжение заголовка вместо маски из звёздочек."""
    word_l = word.lower()
    pd = "персональн" in hint or "152" in hint or "пдн" in hint
    ib = any(x in hint for x in ("информацион", "иб", "конфиденциаль", "коммерческ"))
    labor = any(x in hint for x in ("труд", "тк рф", "работник"))

    mapping = {
        "обработка": "персональных данных" if pd else "информации" if ib else "данных",
        "хранение": "персональных данных" if pd else "информации" if ib else "документов",
        "сбор": "персональных данных" if pd else "информации" if pd or ib else "данных",
        "передача": "персональных данных" if pd else "информации",
        "уничтожение": "персональных данных" if pd else "документов",
        "защита": "персональных данных" if pd else "информации" if ib else "информации",
        "использование": "персональных данных" if pd else "информации",
        "распространение": "персональных данных" if pd else "информации",
        "обезличивание": "персональных данных",
        "блокирование": "персональных данных" if pd else "доступа",
        "трудовые": "отношения" if labor else "отношения",
    }
    return mapping.get(word_l, "")


def _fix_asterisk_mask_runs(text: str, form: dict) -> str:
    """Заменить избыточное маскирование звёздочками в заголовках и терминах."""
    hint = _legal_context_hint(form)
    topic = (form.get("document_topic") or "").strip()

    def replace_heading(match: re.Match) -> str:
        prefix = match.group(1).rstrip()
        tail = prefix.split()
        last_word = tail[-1] if tail else ""
        continuation = _term_after_incomplete_heading(last_word, hint)
        if continuation:
            return f"{prefix} {continuation}"
        if topic and len(topic) < 100:
            return f"{prefix} {topic.lower()}"
        return prefix

    result = re.sub(
        r"(?m)(^[\d\.]+\s+[^\*\n]{2,}?)\s+\*{3,}\s*$",
        replace_heading,
        text,
    )
    result = re.sub(
        r"(?i)\b(Обработка|Хранение|Сбор|Передача|Уничтожение|Защита|Использование)\s+\*{3,}",
        lambda m: f"{m.group(1)} {_term_after_incomplete_heading(m.group(1), hint) or 'данных'}",
        result,
    )
    return result


def _apply_organization_name(text: str, organization_name: str) -> str:
    """Подставить наименование организации вместо маскированных или обобщённых обозначений."""
    if not organization_name:
        return text

    replacements = [
        (r"Организация\s+\*{3,}", organization_name),
        (r"организация\s+\*{3,}", organization_name),
        (r"Наименование\s+организации\s*:?\s*\*{3,}", f"Наименование организации: {organization_name}"),
        (r"Общество\s+\*{3,}", organization_name),
        (r"Компания\s+\*{3,}", organization_name),
        (r"Оператор\s+\*{3,}", f"Оператор {organization_name}"),
        (r"Работодатель\s+\*{3,}", f"Работодатель {organization_name}"),
    ]
    result = text
    for pattern, repl in replacements:
        result = re.sub(pattern, repl, result, flags=re.IGNORECASE)
    return result


def finalize_new_vnd_document(raw_doc: str, form: dict) -> str:
    """Постобработка нового ВНД: термины, организация, умеренное маскирование."""
    from new_vnd_form import build_organization_name

    org_name = form.get("organization_name") or build_organization_name(form.get("ownership_form", ""))
    result = _fix_asterisk_mask_runs(raw_doc, form)
    result = _apply_organization_name(result, org_name)
    result = mask_personal_placeholders(result, aggressive=False)
    return result


def _build_rework_plan(analysis_text: str, main_text: str) -> str:
    """Сформировать детальный план переработки по каждому недостатку из отчёта анализа."""
    if not analysis_text.strip():
        return "Отчёт анализа отсутствует — переработай документ по полному соответствию основному ФЗ."

    client = _get_openai_client()
    prompt = f"""На основе отчёта правового анализа ВНД составь ДЕТАЛЬНЫЙ ПЛАН ПЕРЕРАБОТКИ документа.

ОТЧЁТ АНАЛИЗА:
{analysis_text[:14000]}

ФРАГМЕНТ ИСХОДНОГО ВНД:
{main_text[:5000]}

Извлеки ВСЕ недостатки, несоответствия, отсутствующие статьи ФЗ, неучтённые НПА (ТИП Б: ГОСТ, подзаконные акты), рекомендации и риски.
Для КАЖДОГО пункта укажи в формате:

N. [Статья/пункт ФЗ, ГОСТ или раздел ВНД]
   Проблема: ...
   Действие: добавить / изменить / удалить — Глава X, Статья X.Y «...»; при неучтённом НПА — включить в список используемых документов
   Текст статьи (черновик): преамбула + <<RED>>Необходимо указать самостоятельно ...<<ENDRED>> — если нужны данные организации

Не объединяй пункты. Не пропускай постатейные замечания из отчёта.
Минимум 8 пунктов, если в отчёте столько недостатков."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": "Составляй исчерпывающий план переработки. Только план, без вступлений.",
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
            max_tokens=3500,
        )
        plan = (response.choices[0].message.content or "").strip()
        return plan or "Используй все недостатки из отчёта анализа для переработки."
    except Exception as exc:
        logger.warning("Не удалось составить план переработки: %s", exc)
        return "Устрани все недостатки из отчёта анализа, добавив полный текст статей."


def _resolve_rework_metadata(
    main_filename: str,
    vnd_name: str = "",
    stage1: Optional[dict] = None,
) -> dict:
    """Получить название, сферу и ОПФ для промпта переработки."""
    if stage1:
        from pre_analysis import normalize_stage1_answers

        normalized = normalize_stage1_answers(stage1)
        return {
            "document_name": vnd_name or os.path.basename(main_filename),
            "activity_sphere": normalized.get("activity_sphere", "не указана"),
            "ownership_form": normalized.get("ownership_form", "не указана"),
            "legal_areas": normalized.get("legal_areas") or [],
        }

    detected = detect_rework_stage1(main_filename, vnd_name)
    return {
        "document_name": vnd_name or os.path.basename(main_filename),
        "activity_sphere": detected.get("activity_sphere") or "не указана",
        "ownership_form": detected.get("ownership_form") or "не указана",
        "legal_areas": detected.get("legal_areas") or [],
    }


def generate_rework_document(
    main_filename: str,
    analysis_text: str,
    vnd_name: str = "",
    stage1: Optional[dict] = None,
) -> dict:
    main_text = read_create_file(main_filename)
    meta = _resolve_rework_metadata(main_filename, vnd_name, stage1)
    rework_plan = _build_rework_plan(analysis_text, main_text)
    refs_context, refs_meta = _prepare_federal_refs_context(main_filename, main_text)
    gost_context = get_gost_context()
    fz_query = (
        f"{meta['activity_sphere']} {' '.join(meta.get('legal_areas') or [])} "
        f"{analysis_text[:500]} {main_text[:300]}"
    )
    fz_context = get_fz_context(fz_query, max_results=8)

    legal_areas = ", ".join(meta.get("legal_areas") or []) or "не указаны"

    user_content = f"""=== ИСХОДНЫЕ ДАННЫЕ ===
Наименование документа: {meta['document_name']}
Сфера деятельности: {meta['activity_sphere']}
Организационно-правовая форма: {meta['ownership_form']}
Области законодательства: {legal_areas}

=== ПЛАН ПЕРЕРАБОТКИ (выполни КАЖДЫЙ пункт — отрази в статьях ВНД) ===
{rework_plan[:6000]}

=== ОТЧЁТ АНАЛИЗА (источник недостатков) ===
{analysis_text[:12000]}

ВАЖНО: отчёт анализа обязателен к исполнению. Все рекомендации, недостатки и требования из отчёта должны быть внесены в переработанный ВНД (изменения, дополнения, удаление устаревших норм). Неучтённые в исходном ВНД НПА из отчёта (ФЗ, ГОСТ и др.) — добавь в список используемых документов и отрази в статьях. Для сведений, известных только организации, используй маркер <<RED>>Необходимо указать самостоятельно …<<ENDRED>>.

=== ИСХОДНЫЙ ДОКУМЕНТ (переработать, не копировать без изменений) ===
{main_text[:10000]}

=== ФЕДЕРАЛЬНЫЕ ССЫЛКИ И АКТУАЛЬНОСТЬ ===
{refs_context[:2500]}

=== ГОСТ (оформление и структура) ===
{gost_context[:2500]}

=== ФЕДЕРАЛЬНЫЙ ЗАКОН (адаптировать требования в статьи ВНД) ===
{fz_context[:5000]}

Сформируй ПОЛНЫЙ переработанный ВНД: главы со статьями и развёрнутым текстом каждой статьи.
Каждый пункт плана переработки должен быть отражён в тексте документа."""

    client = _get_openai_client()
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": REWORK_PROMPT},
            {"role": "user", "content": user_content},
        ],
        temperature=0.35,
        max_tokens=16000,
    )
    raw_doc = (response.choices[0].message.content or "").strip()
    document = mask_personal_placeholders(raw_doc)

    return {
        "document": document,
        "refs_meta": refs_meta,
        "gost_used": bool(gost_context),
        "metadata": meta,
        "rework_plan": rework_plan[:2000],
    }


def build_rework_changes_report(
    original_text: str,
    reworked_text: str,
    analysis_text: str,
    document_name: str = "ВНД",
) -> str:
    """Отчёт о переработке по шаблону: удалено / изменено / добавлено."""
    report_date = datetime.now().strftime("%d.%m.%Y")
    system_prompt = REWORK_CHANGES_REPORT_PROMPT.format(
        document_name=document_name,
        report_date=report_date,
    )
    user_content = f"""=== ОТЧЁТ АНАЛИЗА ===
{analysis_text[:10000]}

=== ИСХОДНЫЙ ДОКУМЕНТ ===
{original_text[:8000]}

=== ПЕРЕРАБОТАННЫЙ ДОКУМЕНТ ===
{reworked_text[:10000]}

Составь отчёт о переработке по шаблону."""

    client = _get_openai_client()
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content},
            ],
            temperature=0.2,
            max_tokens=2500,
        )
        report = (response.choices[0].message.content or "").strip()
        if report:
            return report
    except Exception as exc:
        logger.warning("Не удалось сформировать отчёт о переработке: %s", exc)

    return f"""ОТЧЁТ О ПЕРЕРАБОТКЕ ВНД
Документ: {document_name}
Дата: {report_date}

Удалено — статьи:
нет

Изменены — статьи:
• см. переработанный документ — изменения по отчёту анализа

Добавлены — статьи:
• см. переработанный документ — дополнения по отчёту анализа
"""


def save_rework_changes_report(report: str, title: str) -> dict:
    """Сохранить отчёт о переработке в папку new-doc."""
    folder = settings.new_doc_folder
    os.makedirs(folder, exist_ok=True)
    safe_title = re.sub(r'[<>:"/\\|?*]', "_", title or "ВНД")[:80]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Отчёт_переработки_{safe_title}_{timestamp}.txt"
    filepath = os.path.join(folder, filename)
    with open(filepath, "w", encoding="utf-8") as handle:
        handle.write(report)
    return {"filename": filename, "filepath": filepath, "format": "txt", "folder": folder}


def analyze_new_vnd_task(form_data: dict) -> dict:
    """Анализ вводных данных для создания нового ВНД (этап 1)."""
    from new_vnd_form import form_to_prompt_context, normalize_new_vnd_form
    from new_vnd_followup import checklist_to_prompt_text

    form = normalize_new_vnd_form(form_data)
    checklist = checklist_to_prompt_text(form.get("legal_area", "custom"))
    user_content = f"""ВВОДНЫЕ ДАННЫЕ:

{form_to_prompt_context(form)}

ОБЯЗАТЕЛЬНЫЙ ЧЕК-ЛИСТ РАЗДЕЛОВ БУДУЩЕГО ДОКУМЕНТА:
{checklist}

Выполни анализ и сформируй результат в указанном формате."""

    client = _get_openai_client()
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": NEW_VND_ANALYSIS_PROMPT},
            {"role": "user", "content": user_content},
        ],
        temperature=0.4,
        max_tokens=3500,
    )
    analysis_text = (response.choices[0].message.content or "").strip()
    laws = extract_laws_from_analysis(analysis_text, form)

    return {
        "analysis": analysis_text,
        "form": form,
        "laws": laws,
    }


def extract_laws_from_analysis(analysis_text: str, form: dict) -> List[dict]:
    """Извлечь список НПА из результата анализа для скачивания."""
    prompt = f"""По результату правового анализа и вводным данным составь JSON со списком
федеральных законов, кодексов, постановлений и ГОСТ для скачивания.

Вводные:
Название: {form.get('document_name', '')}
Область: {form.get('legal_area_resolved', '')}
Сфера: {form.get('activity_sphere_resolved', '')}

Результат анализа:
{analysis_text[:9000]}

Верни JSON:
{{
  "laws": [
    {{
      "title": "краткое название",
      "number": "152-ФЗ или пусто",
      "search_query": "строка для поиска на pravo.gov.ru"
    }}
  ]
}}
От 3 до 12 наиболее релевантных документов."""

    try:
        data = _llm_json(prompt, max_tokens=1200)
        laws = data.get("laws") or []
        normalized = []
        for law in laws:
            if not isinstance(law, dict):
                continue
            normalized.append({
                "title": law.get("title") or law.get("search_query") or "Федеральный закон",
                "number": law.get("number") or "",
                "search_query": law.get("search_query") or law.get("title") or "",
                "raw": law.get("search_query") or law.get("title") or "",
            })
        if normalized:
            return normalized
    except Exception as exc:
        logger.warning("Не удалось извлечь список НПА из анализа: %s", exc)

    return identify_required_federal_laws({
        "activity_sphere": form.get("activity_sphere_resolved", ""),
        "ownership_form": form.get("ownership_form", ""),
        "legal_areas_text": form.get("legal_area_resolved", ""),
    })


def _build_fz_context_for_laws(laws: List[dict], form: dict) -> str:
    fz_parts = []
    for law in laws[:10]:
        query = law.get("search_query") or law.get("title") or ""
        chunk = get_fz_context(query, max_results=2)
        if chunk:
            fz_parts.append(f"=== {law.get('title', query)} ===\n{chunk}")
    if not fz_parts:
        fz_query = f"{form.get('activity_sphere_resolved', '')} {form.get('legal_area_resolved', '')}"
        fz_parts.append(get_fz_context(fz_query, max_results=6))
    return "\n\n".join(fz_parts)[:12000]


def generate_new_vnd_plan(form: dict, analysis_text: str, checklist: str) -> List[dict]:
    """Сформировать план разделов документа."""
    from new_vnd_followup import get_section_blueprint
    from new_vnd_form import form_to_prompt_context

    user_content = f"""ВВОДНЫЕ:
{form_to_prompt_context(form)}

АНАЛИЗ:
{analysis_text[:5000]}

ЧЕК-ЛИСТ:
{checklist}

БАЗОВЫЙ ШАБЛОН РАЗДЕЛОВ:
{json.dumps(get_section_blueprint(form.get('legal_area', 'custom')), ensure_ascii=False)}"""

    try:
        client = _get_openai_client()
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": NEW_VND_PLAN_PROMPT},
                {"role": "user", "content": user_content},
            ],
            temperature=0.3,
            max_tokens=2000,
        )
        raw = (response.choices[0].message.content or "").strip()
        raw = re.sub(r"^```json\s*|\s*```$", "", raw, flags=re.IGNORECASE)
        data = json.loads(raw)
        sections = data.get("sections") or []
        if sections and isinstance(sections, list):
            return sections
    except Exception as exc:
        logger.warning("Не удалось сформировать план документа: %s", exc)

    return get_section_blueprint(form.get("legal_area", "custom"))


def generate_new_vnd_section(
    form: dict,
    analysis_text: str,
    section: dict,
    checklist: str,
    fz_context: str,
    gost_context: str,
    prior_text: str = "",
) -> str:
    """Сгенерировать один раздел документа."""
    from new_vnd_form import build_organization_name, form_to_prompt_context

    org_display = form.get("organization_name") or build_organization_name(form.get("ownership_form", ""))
    section_title = section.get("title") or section.get("id") or "Раздел"
    requirements = section.get("requirements") or ""
    min_articles = section.get("min_articles") or section.get("articles") or 4

    user_content = f"""=== ОРГАНИЗАЦИЯ ===
{org_display}

=== ВВОДНЫЕ ===
{form_to_prompt_context(form)}

=== РАЗДЕЛ ДЛЯ ГЕНЕРАЦИИ ===
Заголовок: {section_title}
Требования: {requirements}
Минимум статей/пунктов: {min_articles}

=== ЧЕК-ЛИСТ (ориентир) ===
{checklist[:3000]}

=== АНАЛИЗ ===
{analysis_text[:3000]}

=== УЖЕ СГЕНЕРИРОВАННЫЕ РАЗДЕЛЫ (не повторять) ===
{prior_text[-4000:] if prior_text else 'нет'}

=== ФЕДЕРАЛЬНЫЕ НПА ===
{fz_context[:6000]}

=== ГОСТ ===
{gost_context[:2500]}

Напиши только раздел «{section_title}»."""

    client = _get_openai_client()
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": NEW_VND_SECTION_PROMPT},
            {"role": "user", "content": user_content},
        ],
        temperature=0.4,
        max_tokens=4500,
    )
    return (response.choices[0].message.content or "").strip()


def generate_new_vnd_recommendations(
    form: dict,
    analysis_text: str,
    document_text: str,
    fz_context: str = "",
) -> str:
    """Сгенерировать блок «Рекомендации» в конце документа."""
    from new_vnd_form import build_organization_name, form_to_prompt_context
    from new_vnd_followup import companion_docs_hints_to_prompt_text

    org_display = form.get("organization_name") or build_organization_name(form.get("ownership_form", ""))
    legal_area = form.get("legal_area", "custom")
    hints = companion_docs_hints_to_prompt_text(legal_area)

    user_content = f"""=== ОРГАНИЗАЦИЯ ===
{org_display}

=== СОЗДАННЫЙ ДОКУМЕНТ ===
Наименование: {form.get('document_name', '')}

=== ВВОДНЫЕ ===
{form_to_prompt_context(form)}

=== АНАЛИЗ ===
{analysis_text[:2500]}

=== ФРАГМЕНТ СОЗДАННОГО ТЕКСТА (для контекста) ===
{document_text[-6000:] if document_text else 'нет'}

=== ОРИЕНТИРЫ СОПУТСТВУЮЩИХ ДОКУМЕНТОВ ===
{hints}

=== ФЕДЕРАЛЬНЫЕ НПА (кратко) ===
{fz_context[:3000]}

Составь раздел «Рекомендации по разработке дополнительных документов»."""

    client = _get_openai_client()
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": NEW_VND_RECOMMENDATIONS_PROMPT},
            {"role": "user", "content": user_content},
        ],
        temperature=0.35,
        max_tokens=2500,
    )
    return (response.choices[0].message.content or "").strip()


def _assemble_document_header(form: dict) -> str:
    name = form.get("document_name") or "ВНД"
    org = form.get("organization_name") or "Организация"
    return f"{name}\n\n(Проект локального нормативного акта {org})\n\n"


def _assemble_document_footer(form: dict) -> str:
    org = form.get("organization_name") or "Организация"
    year = _current_calendar_year()
    return (
        f"\n\n---\n\nУтверждено:\n"
        f"Генеральный директор {org}\n"
        f"«___» ____________ {year}г.\n"
        f"(подпись)\n\n"
        f"Дата утверждения: {_masked_current_date()}\n"
    )


def generate_new_document_v2(
    form_data: dict,
    analysis_text: str = "",
    laws: Optional[List[dict]] = None,
    download_result: Optional[dict] = None,
) -> dict:
    """Создание нового ВНД: план → разделы → сборка."""
    from new_vnd_form import normalize_new_vnd_form
    from new_vnd_followup import checklist_to_prompt_text

    form = normalize_new_vnd_form(form_data)
    laws = laws or extract_laws_from_analysis(analysis_text, form)

    if download_result is None:
        download_result = download_required_laws(laws)

    gost_context = get_gost_context()
    fz_context = _build_fz_context_for_laws(laws, form)
    checklist = checklist_to_prompt_text(form.get("legal_area", "custom"))

    sections = generate_new_vnd_plan(form, analysis_text, checklist)
    logger.info("План документа: %s разделов", len(sections))

    parts = [_assemble_document_header(form)]
    for index, section in enumerate(sections):
        title = section.get("title") or section.get("id") or f"Раздел {index + 1}"
        logger.info("Генерация раздела %s/%s: %s", index + 1, len(sections), title)
        section_text = generate_new_vnd_section(
            form=form,
            analysis_text=analysis_text,
            section=section,
            checklist=checklist,
            fz_context=fz_context,
            gost_context=gost_context,
            prior_text="\n\n".join(parts),
        )
        if section_text:
            parts.append(section_text)

    main_body = "\n\n".join(parts)
    parts.append(_assemble_document_footer(form))

    logger.info("Генерация блока «Рекомендации»...")
    recommendations = generate_new_vnd_recommendations(
        form=form,
        analysis_text=analysis_text,
        document_text=main_body,
        fz_context=fz_context,
    )
    if recommendations:
        parts.append(recommendations)

    raw_doc = "\n\n".join(parts)
    document = finalize_new_vnd_document(raw_doc, form)

    return {
        "document": document,
        "form": form,
        "analysis": analysis_text[:2000],
        "laws": laws,
        "download_result": download_result,
        "gost_used": bool(gost_context),
        "plan_sections": [s.get("title") or s.get("id") for s in sections],
        "generation_mode": "multi_section",
        "has_recommendations": bool(recommendations),
    }


def generate_new_document(stage1: dict) -> dict:
    legal_text = stage1.get("legal_areas_text") or ", ".join(stage1.get("legal_areas") or [])
    gost_context = get_gost_context()
    fz_query = f"{stage1.get('activity_sphere', '')} {legal_text}"
    fz_context = get_fz_context(fz_query, max_results=5)

    user_content = f"""=== ПАРАМЕТРЫ ОРГАНИЗАЦИИ ===
Сфера деятельности: {stage1.get('activity_sphere', '')}
Форма собственности: {stage1.get('ownership_form', '')}
Требования по законодательству: {legal_text}

=== ГОСТ Р 7.0.97-2025 ===
{gost_context[:2500]}

=== ФЕДЕРАЛЬНОЕ ЗАКОНОДАТЕЛЬСТВО ===
{fz_context[:3000]}

Подготовь проект нового ВНД по указанным параметрам."""

    client = _get_openai_client()
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": NEW_VND_PROMPT},
            {"role": "user", "content": user_content},
        ],
        temperature=0.5,
        max_tokens=4000,
    )
    raw_doc = (response.choices[0].message.content or "").strip()
    document = mask_personal_placeholders(raw_doc)

    return {
        "document": document,
        "stage1": stage1,
        "gost_used": bool(gost_context),
    }


def _add_docx_paragraph_with_red_markers(doc: Document, text: str) -> None:
    """Добавить абзац в docx с жирным красным текстом внутри <<RED>>…<<ENDRED>>."""
    paragraph = doc.add_paragraph()
    pos = 0
    for match in RED_MARKER_PATTERN.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        run = paragraph.add_run(match.group(1))
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build_generated_document(
    document: str,
    title: str,
    fmt: str = "txt",
) -> tuple[bytes, str, str]:
    """Сформировать документ ВНД в памяти. Возвращает (bytes, filename, media_type)."""
    safe_title = re.sub(r'[<>:"/\\|?*]', "_", title or "ВНД")[:80]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    if fmt == "docx":
        filename = f"ВНД_{safe_title}_{timestamp}.docx"
        buffer = io.BytesIO()
        doc = Document()
        for paragraph in document.split("\n"):
            if "<<RED>>" in paragraph:
                _add_docx_paragraph_with_red_markers(doc, paragraph)
            else:
                doc.add_paragraph(paragraph)
        doc.save(buffer)
        return (
            buffer.getvalue(),
            filename,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    filename = f"ВНД_{safe_title}_{timestamp}.txt"
    return document.encode("utf-8"), filename, "text/plain; charset=utf-8"


def save_generated_document(
    document: str,
    title: str,
    fmt: str = "txt",
    output_folder: Optional[str] = None,
) -> dict:
    folder = output_folder or settings.out_folder
    os.makedirs(folder, exist_ok=True)
    content, filename, _media_type = build_generated_document(document, title, fmt)
    filepath = os.path.join(folder, filename)
    with open(filepath, "wb") as handle:
        handle.write(content)

    return {"filename": filename, "filepath": filepath, "format": fmt, "folder": folder}


def build_create_qa_system_prompt(mode: str, document_text: str, title: str) -> str:
    """Системный промпт: база создания ВНД + режим ответов на вопросы."""
    base = REWORK_PROMPT if (mode or "").lower() == "rework" else NEW_VND_PROMPT
    doc_excerpt = (document_text or "")[:14000]
    safe_title = (title or "ВНД").strip()
    return f"""{base}

{CREATE_QA_MODE_SUFFIX}

=== ПОДГОТОВЛЕННЫЙ ВНД («{safe_title}») ===
{doc_excerpt}
"""


def answer_create_document_question(
    mode: str,
    document_text: str,
    title: str,
    messages: List[dict],
    user_message: str,
) -> dict:
    """Ответ на вопрос по подготовленному ВНД с учётом истории диалога."""
    user_message = (user_message or "").strip()
    if not user_message:
        raise ValueError("Введите вопрос")
    if not (document_text or "").strip():
        raise ValueError("Текст документа отсутствует")

    history: List[dict] = []
    for item in messages or []:
        role = item.get("role")
        content = (item.get("content") or "").strip()
        if role in ("user", "assistant") and content:
            history.append({"role": role, "content": content})

    system_prompt = build_create_qa_system_prompt(mode, document_text, title)
    api_messages: List[dict] = [{"role": "system", "content": system_prompt}]
    api_messages.extend(history[-24:])
    api_messages.append({"role": "user", "content": user_message})

    client = _get_openai_client()
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=api_messages,
        temperature=0.4,
        max_tokens=2000,
    )
    reply = (response.choices[0].message.content or "").strip()
    if not reply:
        reply = "Не удалось сформировать ответ. Попробуйте переформулировать вопрос по документу."

    updated = history + [
        {"role": "user", "content": user_message},
        {"role": "assistant", "content": reply},
    ]
    return {"reply": reply, "messages": updated}


def build_create_qa_dialog(messages: List[dict], title: str) -> tuple[bytes, str, str]:
    """Сформировать файл диалога в памяти."""
    if not messages:
        raise ValueError("Диалог пуст — нет сообщений для сохранения")

    safe_title = re.sub(r'[<>:"/\\|?*]', "_", title or "ВНД")[:80]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Диалог_{safe_title}_{timestamp}.txt"

    lines = [
        "ДИАЛОГ ПО ВНД",
        f"Документ: {title or 'ВНД'}",
        f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
        "",
    ]
    for item in messages:
        role = item.get("role")
        content = (item.get("content") or "").strip()
        if not content:
            continue
        label = "Пользователь" if role == "user" else "Консультант"
        lines.append(f"{label}:")
        lines.append(content)
        lines.append("")

    return "\n".join(lines).strip().encode("utf-8") + b"\n", filename, "text/plain; charset=utf-8"


def save_create_qa_dialog(messages: List[dict], title: str, mode: str) -> dict:
    """Сохранить диалог по ВНД рядом с основным документом."""
    folder = settings.new_doc_folder if (mode or "").lower() == "rework" else settings.out_folder
    os.makedirs(folder, exist_ok=True)
    content, filename, _media_type = build_create_qa_dialog(messages, title)
    filepath = os.path.join(folder, filename)
    with open(filepath, "wb") as handle:
        handle.write(content)

    return {"filename": filename, "filepath": filepath, "format": "txt", "folder": folder}


def generate_rework_from_analysis(
    main_filename: str,
    analysis_text: str,
    vnd_name: str = "",
    stage1: Optional[dict] = None,
    analysis_meta: Optional[dict] = None,
) -> dict:
    """Сгенерировать переработанный ВНД по готовому тексту анализа."""
    if not (analysis_text or "").strip():
        raise ValueError("Текст отчёта анализа отсутствует")

    effective_stage1 = stage1
    gen = generate_rework_document(
        main_filename,
        analysis_text,
        vnd_name,
        stage1=effective_stage1,
    )
    gen["analysis_text"] = analysis_text[:2000]
    gen["analysis_meta"] = analysis_meta or {}

    doc_title = vnd_name or gen.get("metadata", {}).get("document_name") or "Переработка"
    original_text = read_create_file(main_filename)
    try:
        changes_report = build_rework_changes_report(
            original_text,
            gen["document"],
            analysis_text,
            doc_title,
        )
        gen["changes_report"] = changes_report
    except Exception as exc:
        logger.warning("Не удалось подготовить отчёт о переработке: %s", exc)
        gen["changes_report"] = ""

    return gen


def process_rework(
    main_filename: str,
    analysis_filename: Optional[str] = None,
    vnd_name: str = "",
    stage1: Optional[dict] = None,
    analysis_text: Optional[str] = None,
) -> dict:
    analysis_meta: dict = {}
    effective_stage1 = stage1

    if (analysis_text or "").strip():
        text = analysis_text.strip()
        analysis_meta = {"from_text": True, "saved_to_file": False}
    elif analysis_filename:
        from search_vnd import validate_analysis_for_main

        validate_analysis_for_main(main_filename, analysis_filename)
        text = read_analysis_report_file(analysis_filename)
        analysis_meta = {"uploaded_report": True, "filename": analysis_filename}
    else:
        if not stage1:
            raise ValueError("Для анализа без отчёта необходимы параметры этапа 1")
        analyzed = analyze_for_rework(main_filename, vnd_name, stage1)
        text = analyzed["analysis_text"]
        effective_stage1 = analyzed["stage1"]
        analysis_meta = analyzed["analysis_meta"]

    return generate_rework_from_analysis(
        main_filename,
        text,
        vnd_name,
        stage1=effective_stage1,
        analysis_meta=analysis_meta,
    )
